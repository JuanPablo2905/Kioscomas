from pathlib import Path
from docx import Document
from docx.shared import Pt


path = Path("DOCUMENTACION_KIOSCOAPP.docx")
doc = Document(path)

marker = "Avances técnicos completados — 17 de julio de 2026"
if not any(marker in paragraph.text for paragraph in doc.paragraphs):
    doc.add_page_break()
    doc.add_heading(marker, level=1)
    doc.add_paragraph(
        "Actualización del estado técnico posterior a la migración del monolito "
        "a un proyecto local y a la generación de la primera versión portable para Windows."
    )

    completed = [
        "Crear un proyecto local compilable con Vite, React y Tailwind.",
        "Separar las pantallas por menú: Inicio, Stock, Vitrina, Ventas/Caja, Compras, Clientes/Fiado, Reportes y Administración.",
        "Separar Autenticación, layout, datos semilla y utilidades compartidas.",
        "Conservar KioscoAppv4.jsx como respaldo del monolito original.",
        "Implementar persistencia compatible con window.storage de Claude y localStorage en navegador/Electron.",
        "Proteger la carga ante APIs de almacenamiento incompatibles o errores de acceso.",
        "Configurar rutas relativas de Vite para ejecutar recursos desde un archivo portable.",
        "Integrar Electron y electron-builder.",
        "Generar release/KioscoApp-0.1.0-portable.exe, sin instalación.",
        "Corregir la pantalla blanca de la primera versión portable.",
        "Corregir la dependencia de permisos de la pantalla Inicio.",
        "Agregar una prueba automática de arranque que confirma que aparece el login.",
        "Verificar la compilación de producción con Vite.",
    ]
    for item in completed:
        doc.add_paragraph(f"☑ {item}", style="List Bullet")

    doc.add_heading("Pendientes relacionados", level=2)
    pending = [
        "Migrar de Electron a Tauri si se prioriza reducir el peso del ejecutable.",
        "Reemplazar localStorage por IndexedDB/Dexie.",
        "Agregar ícono y metadatos definitivos del producto.",
        "Firmar digitalmente el ejecutable para evitar advertencias de SmartScreen.",
        "Definir el mecanismo de actualización y distribución.",
        "Extraer hooks y subdividir los módulos más grandes, especialmente Ventas y Administración.",
    ]
    for item in pending:
        doc.add_paragraph(f"☐ {item}", style="List Bullet")

    for style_name in ("Normal", "List Bullet"):
        style = doc.styles[style_name]
        style.font.name = "Arial"
        style.font.size = Pt(10.5)

    doc.save(path)

print(path.resolve())
